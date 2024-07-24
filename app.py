 # -*- coding: utf-8 -*-
import requests
from requests.exceptions import HTTPError
import json
import os
import sys
import time
from flask import Flask, request, render_template_string
import subprocess
from docx import Document
from docx.shared import Inches, Pt

app = Flask(__name__)

@app.route('/')
def index():
    return render_template_string(open('index.html').read())

@app.route('/submit', methods=['POST'])

#define what happens once user click on submit
def submit():
    user_seqdia_text = request.form['user_seqdia_text']
    user_netdia_text = request.form['user_netdia_text']
    interface_name = request.form['interface_name']
    canonical_json= request.form['canonical_json']

    # print(user_seqdia_text)
    # print(user_netdia_text)

    #define OpenAI API url
    url = "https://bot-svc-llm.sfproxy.einstein.aws-dev4-uswest2.aws.sfdc.cl/v1.0/generations"

    # Calling the functions one by one
    sequenceDiagramText = generateSeqDiagram(user_seqdia_text,url)
    time.sleep(1)  # Makes Python wait for 1 second
    netDiagramText= generateNetDiagram(user_netdia_text,url)

    sequenceDiagramText= sequenceDiagramText.replace("(", "_").replace(")", "")
    netDiagramText= netDiagramText.replace("(", "_").replace(")", "")

    
    ###########################################################################################

    #  # Save the original stdout
    # original_stdout = sys.stdout


    f = open("sequenceDiagramUmlText.txt", "a")
    f.write(sequenceDiagramText)
    f.close()


    f = open("netDiagramText.txt", "a")
    f.write(netDiagramText)
    f.close()

    # time.sleep(2)

    with open("sequenceDiagramUmlText.txt", "r+") as output:
      subprocess.call(["python", "-m","plantuml", "sequenceDiagramUmlText.txt"], stdout=output);
      print('seq diagram generated')

    with open("netDiagramText.txt", "r+") as output:
      subprocess.call(["python", "-m","plantuml", "netDiagramText.txt"], stdout=output);
      print('network diagram generated')
  
    
    ###########################################################################################


     # call more functions
    apiTable=generateAPITable(user_seqdia_text,url)
    apiProtocols=generateAPIProtocols(user_seqdia_text,url)
    errorHandlingVerbiage=generateErrorHandling(user_seqdia_text,url)
    loggingVerbiage=generateLogging(user_seqdia_text,url)
    mappingVerbiage=generateBasicMapping(canonical_json,url)


    # create DOC file
    docTitle= 'Mulesoft Technical Specification: ' + interface_name
    docName= interface_name + '.docx'
    docFileCreationResult= craeteDocx(docTitle,docName,user_seqdia_text,apiTable,apiProtocols,mappingVerbiage,loggingVerbiage,errorHandlingVerbiage,user_netdia_text)
    print(docFileCreationResult)
 
    returnMessage= "Thank you! your Spec v1 generated. Name of the file is: " +  docName + "\r\nNote: This Spec is AI generated and may not be accurate. Please check the content and edit as needed. \r\nThank you for using the app. Hope you liked it! "

    return returnMessage

####### Function Declaration Starts #######

def generateSeqDiagram(seqDiagramPrompt,url):

  seqDiagramPrompt= "Provide a UML texual representation of a sequence diagram for plantuml based on the following steps ? Replace all dashes in all the names with underscores. Do not retun any extra verbiages. Escape ASCII unicode chars as needed. " + seqDiagramPrompt

  # replace dashed with underscores
  seqDiagramPrompt= seqDiagramPrompt.replace("-", "_")

  body = {
    "prompt": seqDiagramPrompt,
    "model": "llmgateway__OpenAIGPT35Turbo"
  }

  print(body)
  
  headers = {
    'Content-Type': 'application/json',
    'Accept': 'application/json',
    'Authorization': 'API_KEY 651192c5-37ff-440a-b930-7444c69f4422',
    'x-client-feature-id': 'EinsteinDocsAnswers',
    'x-sfdc-app-context': 'EinsteinGPT',
    'x-sfdc-core-tenant-id': 'core/prod1/00DDu0000008cuqMAA'
  }

  

  resp = requests.post(
      url,
      json=body,
      headers=headers,
      verify=False
      )
  
  print(resp)

  # access JSON content
  jsonResponse = resp.json()
  print("Entire JSON response")
  print(jsonResponse)

  print(type(jsonResponse))
  print(jsonResponse["id"])
 
  sequenceDiagramText = jsonResponse["generations"][0]["text"]
  print(sequenceDiagramText)
  
  print("successfully generated API sequence Diagram Verbiage")
  return sequenceDiagramText

def generateNetDiagram(netDiagramPrompt,url):
   # append some directions to user content
  netDiagramPrompt = "Provide me a UML texual representation of a network diagram for plantUml for the following cloud architecture. " + netDiagramPrompt + " Escape the front slashes of the CIDRs. Form the actor names like the name provided concatenating with the CIDR. Make sure to Replace all spaces from all names with underscores. Use allowmixing directive but properly put the components inside a box. do NOT mess it up. Put skinparam linetype ortho. Use Cloud, Frame, Rectangle and nodes as needed.  Escape ASCII unicode chars as needed. "

  netDiagramPrompt= netDiagramPrompt.replace("-", "_")


  body = {
    "prompt": netDiagramPrompt,
    "model": "llmgateway__OpenAIGPT35Turbo"
}

  print(body)
  
  headers = {
    'Content-Type': 'application/json',
    'Accept': 'application/json',
    'Authorization': 'API_KEY 651192c5-37ff-440a-b930-7444c69f4422',
    'x-client-feature-id': 'EinsteinDocsAnswers',
    'x-sfdc-app-context': 'EinsteinGPT',
    'x-sfdc-core-tenant-id': 'core/prod1/00DDu0000008cuqMAA'
  }

  # url = "https://bot-svc-llm.sfproxy.einstein.aws-dev4-uswest2.aws.sfdc.cl/v1.0/generations"
  # url = "https://bot-svc-llm.sfproxy.einsteintest1.test1-uswest2.aws.sfdc.cl/v1.0/generations"

  resp = requests.post(
      url,
      json=body,
      headers=headers,
      verify=False
      )
  
  print(resp)
  # resp.raise_for_status()


  # access JSON content
  jsonResponse = resp.json()
  print("Entire JSON response")
  print(jsonResponse)

  print(type(jsonResponse))
  print(jsonResponse["id"])
  
  netDiagramText = jsonResponse["generations"][0]["text"]
  print(netDiagramText)

  print("successfully generated Network Diagram Verbiage")
  return netDiagramText

def generateAPITable(seqDiagramPrompt,url):
   seqDiagramPrompt= "Can you list down the API names only from the verbiage below? Can you also form a table and put the API names is there along with the type. if you see sapi in the name then API type is SYSTEM, if papi then PROCESS and if eapi then EXPERIENCE. Also put the github link in the table. The githunb link is like http://github.com/my-comp/{api-name]. " + seqDiagramPrompt


   body = {
    "prompt": seqDiagramPrompt,
    "model": "llmgateway__OpenAIGPT35Turbo"
  }

   print(body)
  
   headers = {
    'Content-Type': 'application/json',
    'Accept': 'application/json',
    'Authorization': 'API_KEY 651192c5-37ff-440a-b930-7444c69f4422',
    'x-client-feature-id': 'EinsteinDocsAnswers',
    'x-sfdc-app-context': 'EinsteinGPT',
    'x-sfdc-core-tenant-id': 'core/prod1/00DDu0000008cuqMAA'
  }

  

   resp = requests.post(
      url,
      json=body,
      headers=headers,
      verify=False
      )
  
   print(resp)

    # access JSON content
   jsonResponse = resp.json()
   print("Entire JSON response")
   print(jsonResponse)

   print(type(jsonResponse))
   print(jsonResponse["id"])
 
   apiTable = jsonResponse["generations"][0]["text"]
   print(apiTable)
  
   print("successfully generated API Table")
   return apiTable

def generateAPIProtocols(seqDiagramPrompt,url):
   seqDiagramPrompt= "Can you list down the tranmission protocols used for the APIs from the below verbiages.If you see eapi that would be called via OAuth and Azure AD is the Oauth provider. for sapi and papi, the protocols is client authentication and Azure id is the provider. Put everything in a table. Put an extra column in the end and describes each protocol what does it mean. Do NOT enlist endpoint names, only capture api names. " + seqDiagramPrompt


   body = {
    "prompt": seqDiagramPrompt,
    "model": "llmgateway__OpenAIGPT35Turbo"
  }

   print(body)
  
   headers = {
    'Content-Type': 'application/json',
    'Accept': 'application/json',
    'Authorization': 'API_KEY 651192c5-37ff-440a-b930-7444c69f4422',
    'x-client-feature-id': 'EinsteinDocsAnswers',
    'x-sfdc-app-context': 'EinsteinGPT',
    'x-sfdc-core-tenant-id': 'core/prod1/00DDu0000008cuqMAA'
  }

  

   resp = requests.post(
      url,
      json=body,
      headers=headers,
      verify=False
      )
  
   print(resp)

    # access JSON content
   jsonResponse = resp.json()
   print("Entire JSON response")
   print(jsonResponse)

   print(type(jsonResponse))
   print(jsonResponse["id"])
 
   apiProtocols = jsonResponse["generations"][0]["text"]
   print(apiProtocols)
  
   print("successfully generated API Protocols")
   return apiProtocols

def generateErrorHandling(seqDiagramPrompt,url):
   seqDiagramPrompt= "Can you generate verbiages for all probable API errors and error handling mechanism for each API in the below verbiages. Put everything in a table. outside the table generate common notes for HTTP errors and mention that Mulesoft will use a common error handler which is documented at https://sharepoint.mycomp.com/integration-docs/c4eassets/error-handling. " + seqDiagramPrompt


   body = {
    "prompt": seqDiagramPrompt,
    "model": "llmgateway__OpenAIGPT35Turbo"
  }

   print(body)
  
   headers = {
    'Content-Type': 'application/json',
    'Accept': 'application/json',
    'Authorization': 'API_KEY 651192c5-37ff-440a-b930-7444c69f4422',
    'x-client-feature-id': 'EinsteinDocsAnswers',
    'x-sfdc-app-context': 'EinsteinGPT',
    'x-sfdc-core-tenant-id': 'core/prod1/00DDu0000008cuqMAA'
  }

  

   resp = requests.post(
      url,
      json=body,
      headers=headers,
      verify=False
      )
  
   print(resp)

    # access JSON content
   jsonResponse = resp.json()
   print("Entire JSON response")
   print(jsonResponse)

   print(type(jsonResponse))
   print(jsonResponse["id"])
 
   errorHandlingVerbiage = jsonResponse["generations"][0]["text"]
   print(errorHandlingVerbiage)
  
   print("successfully generated ErrorHandling Verbiage")
   return errorHandlingVerbiage

def generateLogging(seqDiagramPrompt,url):
   seqDiagramPrompt= "Can you generate verbiages for possible logging levels like INFO, DEBUG etc. for each API step in a table format? do not return the processing steps again. only API names. additionally return a sample json logger for any step inside an API with different logging attributes like message, timestamp etc. for this below processing steps? do not log sensitive info. " + seqDiagramPrompt


   body = {
    "prompt": seqDiagramPrompt,
    "model": "llmgateway__OpenAIGPT35Turbo"
  }

   print(body)
  
   headers = {
    'Content-Type': 'application/json',
    'Accept': 'application/json',
    'Authorization': 'API_KEY 651192c5-37ff-440a-b930-7444c69f4422',
    'x-client-feature-id': 'EinsteinDocsAnswers',
    'x-sfdc-app-context': 'EinsteinGPT',
    'x-sfdc-core-tenant-id': 'core/prod1/00DDu0000008cuqMAA'
  }

  

   resp = requests.post(
      url,
      json=body,
      headers=headers,
      verify=False
      )
  
   print(resp)

    # access JSON content
   jsonResponse = resp.json()
   print("Entire JSON response")
   print(jsonResponse)

   print(type(jsonResponse))
   print(jsonResponse["id"])
 
   loggingVerbiage = jsonResponse["generations"][0]["text"]
   print(loggingVerbiage)
  
   print("successfully generated API Protocols")
   return loggingVerbiage

def generateBasicMapping(canonical_json,url):
   canonical_json= "Can you predict sample mapping in a table format between eapi, papi and sapi for this below processing steps? Below sample json is for PAPI Canonical where fields are in camel case. EAPI has the same fields but in snake case and sapi has the same Salesfoce naming convention and date formats are yyyyMMdd. Additionally can you generate some sample dataweave script to transform EAPI to PAPI and PAPI to SAPI. Canonical JSON: " + canonical_json


   body = {
    "prompt": canonical_json,
    "model": "llmgateway__OpenAIGPT35Turbo"
  }

   print(body)
  
   headers = {
    'Content-Type': 'application/json',
    'Accept': 'application/json',
    'Authorization': 'API_KEY 651192c5-37ff-440a-b930-7444c69f4422',
    'x-client-feature-id': 'EinsteinDocsAnswers',
    'x-sfdc-app-context': 'EinsteinGPT',
    'x-sfdc-core-tenant-id': 'core/prod1/00DDu0000008cuqMAA'
  }

  

   resp = requests.post(
      url,
      json=body,
      headers=headers,
      verify=False
      )
  
   print(resp)

    # access JSON content
   jsonResponse = resp.json()
   print("Entire JSON response")
   print(jsonResponse)

   print(type(jsonResponse))
   print(jsonResponse["id"])
 
   mappingVerbiage = jsonResponse["generations"][0]["text"]
   print(mappingVerbiage)
  
   print("successfully generated Mapping Verbiage")
   return mappingVerbiage

def craeteDocx(docTitle,docName,user_seqdia_text,apiTable,apiProtocols,mappingVerbiage,loggingVerbiage,errorHandlingVerbiage,user_netdia_text):
   
   
   # Create an instance of a word document
    document = Document()

    #set paragraph styles and font size for table formats
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(7)

#     document.add_paragraph(
#     'first item in unordered list', style='List Bullet'
# )
      # document.add_picture('monty-truth.png', width=Inches(1.25))

    # Add a Title to the document
    document.add_heading(docTitle, 0)

    document.add_heading('Processing Steps')
    run=document.add_paragraph(user_seqdia_text).add_run()
    font=run.font
    run.add_picture('./sequenceDiagramUmlText.png')
    
    document.add_heading('Mulesfot API Details')
    run2=document.add_paragraph(apiTable.replace("\n","\r\n")).add_run()
    font=run2.font

    document.add_heading('API Protocols')
    document.add_paragraph(apiProtocols.replace("\n","\r\n")).add_run()
    font=run.font

    document.add_heading('API Networks and Security')
    document.add_paragraph(user_netdia_text).add_run()
    font=run.font
    run.add_picture('./netDiagramText.png')

    document.add_heading('Data Mapping')
    document.add_paragraph(mappingVerbiage.replace("\n","\r\n")).add_run()
    font=run.font

    document.add_heading('API Logging')
    document.add_paragraph(loggingVerbiage.replace("\n","\r\n")).add_run()
    font=run.font

    document.add_heading('API Error Handling')
    document.add_paragraph(errorHandlingVerbiage.replace("\n","\r\n")).add_run()
    font=run.font

    document.save(docName)

    return "doc file created"


####### Function Declaration Ends #######





if __name__ == '__main__':
    app.run(debug=True)
